import re
import unicodedata
from pathlib import Path

from pypdf import PdfReader
from docx import Document


MIN_PARAGRAPHS = 5
MIN_WORDS = 300

HEADING_WORDS = {"conclusion", "bibliografia", "referencias"}
TITLE_LABELS = {"titulo", "tema"}
MIN_PARAGRAPH_CHARS = 120  # evita contar líneas cortas como referencias


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip().lower()


def _strip_accents(text: str) -> str:
    text = unicodedata.normalize("NFD", text or "")
    return "".join(ch for ch in text if unicodedata.category(ch) != "Mn")


def _clean_text_for_match(text: str) -> str:
    text = (text or "").replace("\ufeff", "").replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    text = _strip_accents(text).lower()
    return text


def _extract_text(file_path: str) -> tuple[str, str]:
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        reader = PdfReader(file_path)
        pages_text = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(pages_text).strip(), "pdf"

    if ext == ".docx":
        doc = Document(file_path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n".join(lines).strip(), "docx"

    return "", "unsupported"


def _extract_docx_paragraphs(file_path: str) -> list[str]:
    doc = Document(file_path)
    out = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            out.append(t)
    return out


def _is_title_label_line(cleaned_line: str) -> bool:
    # "titulo:", "titulo (primera linea):", "tema:"
    return bool(re.match(r"^(titulo|tema)\s*(\([^)]{0,60}\))?\s*:\s*$", cleaned_line))


def _is_title_value_line(cleaned_line: str) -> bool:
    # línea de texto que sigue a "Título (primera línea):"
    words = re.findall(r"\b\w+\b", cleaned_line, flags=re.UNICODE)
    return 3 <= len(words) <= 20 and len(cleaned_line) <= 160


def _detect_title(raw_text: str) -> bool:
    cleaned = _clean_text_for_match(raw_text)
    head = cleaned[:700]

    # Caso 1: inline "Título: Algo"
    inline = re.compile(
        r"(?:^|\A)\s*(titulo|tema)\s*(\([^)]{0,60}\))?\s*:\s*\S+",
        flags=re.IGNORECASE,
    )
    if inline.search(head):
        return True

    return False


def _is_reference_like(cleaned_line: str) -> bool:
    # Detecta líneas de bibliografía típicas
    if re.search(r"\(\d{4}\)", cleaned_line):
        return True
    if "publishing" in cleaned_line:
        return True
    if re.search(r"\bdoi\b|http|www\.", cleaned_line):
        return True
    return False


def _count_paragraphs_docx(file_path: str) -> int:
    paragraphs = _extract_docx_paragraphs(file_path)
    if not paragraphs:
        return 0

    cleaned = [_clean_text_for_match(p) for p in paragraphs]

    count = 0
    i = 0
    while i < len(cleaned):
        line = cleaned[i]
        if not line:
            i += 1
            continue

        # Saltar encabezados
        if line in HEADING_WORDS:
            i += 1
            continue

        # Saltar línea etiqueta de título y su valor inmediato
        if _is_title_label_line(line):
            i += 1
            if i < len(cleaned) and _is_title_value_line(cleaned[i]):
                i += 1
            continue

        # Saltar "bibliografia ...." en la misma línea
        if line.startswith("bibliografia ") or line.startswith("referencias "):
            i += 1
            continue

        # Saltar referencias
        if _is_reference_like(line):
            i += 1
            continue

        # Contar solo párrafos de contenido real
        if len(line) >= MIN_PARAGRAPH_CHARS:
            count += 1

        i += 1

    return min(count, 20)


def _count_paragraphs_pdf(raw_text: str) -> int:
    text = (raw_text or "").strip()
    if not text:
        return 0

    blocks = re.split(r"\n\s*\n", text)
    blocks = [b.strip() for b in blocks if b.strip()]

    count = 0
    for b in blocks:
        c = _clean_text_for_match(b)
        if c in HEADING_WORDS:
            continue
        if c.startswith("bibliografia ") or c.startswith("referencias "):
            continue
        if _is_reference_like(c):
            continue
        if len(c) >= MIN_PARAGRAPH_CHARS:
            count += 1

    if count >= 2:
        return min(count, 20)

    words = re.findall(r"\b\w+\b", text, flags=re.UNICODE)
    return min(max(1, round(len(words) / 110)), 20)


def run_pdf_precheck(file_path: str) -> dict:
    raw_text, source_type = _extract_text(file_path)

    if source_type == "unsupported" or not raw_text:
        return {
            "has_title": False,
            "has_conclusion": False,
            "has_bibliography": False,
            "paragraph_count": 0,
            "word_count": 0,
            "score": 0,
            "passed": False,
            "feedback": "❌ Formato no soportado para precheck. Usa PDF o DOCX.",
        }

    norm = _normalize(raw_text)

    has_title = _detect_title(raw_text)
    has_conclusion = ("conclusión" in norm) or ("conclusion" in norm)
    has_bibliography = (
        ("bibliografía" in norm) or ("bibliografia" in norm) or ("referencias" in norm)
    )

    if source_type == "docx":
        paragraph_count = _count_paragraphs_docx(file_path)
    else:
        paragraph_count = _count_paragraphs_pdf(raw_text)

    word_count = len(re.findall(r"\b\w+\b", raw_text, flags=re.UNICODE))

    score = 0
    score += 25 if has_title else 0
    score += 25 if has_conclusion else 0
    score += 25 if has_bibliography else 0
    score += 25 if paragraph_count >= MIN_PARAGRAPHS else 0

    passed = (
        has_title
        and has_conclusion
        and has_bibliography
        and paragraph_count >= MIN_PARAGRAPHS
        and word_count >= MIN_WORDS
    )

    feedback_items = [
        f"📄 Fuente detectada: {source_type.upper()}",
        "✅ Título detectado." if has_title else "❌ No se detectó título claro.",
        "✅ Conclusión detectada." if has_conclusion else "❌ Falta sección de conclusión.",
        "✅ Bibliografía/referencias detectada."
        if has_bibliography
        else "❌ Falta bibliografía o referencias.",
        f"✅ Párrafos: {paragraph_count} (mínimo {MIN_PARAGRAPHS})."
        if paragraph_count >= MIN_PARAGRAPHS
        else f"❌ Párrafos: {paragraph_count} (mínimo {MIN_PARAGRAPHS}).",
        f"✅ Palabras: {word_count} (mínimo {MIN_WORDS})."
        if word_count >= MIN_WORDS
        else f"❌ Palabras: {word_count} (mínimo {MIN_WORDS}).",
    ]

    return {
        "has_title": has_title,
        "has_conclusion": has_conclusion,
        "has_bibliography": has_bibliography,
        "paragraph_count": paragraph_count,
        "word_count": word_count,
        "score": score,
        "passed": passed,
        "feedback": "\n".join(feedback_items),
    }